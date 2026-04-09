"""Compare two Word documents and generate a revision-tracked output."""

from __future__ import annotations

import difflib

from docx import Document
from docx_revisions import RevisionDocument, RevisionParagraph

from .revision_writer import _enable_markup_view


def _get_paragraph_texts(doc_path: str) -> list[str]:
    """Extract all paragraph texts from a docx."""
    doc = Document(doc_path)
    return [p.text for p in doc.paragraphs]


def diff_documents(
    old_path: str,
    new_path: str,
    author: str = "AI Reviewer",
) -> RevisionDocument:
    """Compare old and new docx files, producing a RevisionDocument with tracked changes.

    Strategy:
    1. Align paragraphs using SequenceMatcher.
    2. For matched paragraphs with differences, use word-level diff to create
       fine-grained tracked deletions + insertions.
    3. Entirely deleted/added paragraphs become block-level tracked changes.
    """
    old_texts = _get_paragraph_texts(old_path)
    new_texts = _get_paragraph_texts(new_path)

    # Work on a copy of the old document to preserve formatting
    rdoc = RevisionDocument(old_path)
    paras = rdoc.paragraphs

    # Use SequenceMatcher to align paragraphs
    sm = difflib.SequenceMatcher(None, old_texts, new_texts)
    opcodes = sm.get_opcodes()

    # We need to process from end to beginning so that index shifts don't affect us
    # when we insert/delete paragraphs. But for same-paragraph replacements we can
    # process in any order since we only modify text within existing paragraphs.

    # First pass: handle "replace" and "equal" at paragraph text level
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            continue
        elif tag == "replace":
            # For each pair of old/new paragraphs, do word-level replacement
            old_chunk = old_texts[i1:i2]
            new_chunk = new_texts[j1:j2]

            # Match up paragraphs 1:1 as much as possible
            pairs = min(len(old_chunk), len(new_chunk))
            for k in range(pairs):
                old_para_idx = i1 + k
                if old_para_idx >= len(paras):
                    break
                _apply_word_level_diff(
                    paras[old_para_idx],
                    old_chunk[k],
                    new_chunk[k],
                    author,
                )

            # Extra new paragraphs → tracked insertions at last matched para
            if len(new_chunk) > len(old_chunk):
                anchor_idx = min(i2 - 1, len(paras) - 1)
                if anchor_idx >= 0:
                    for extra_text in new_chunk[pairs:]:
                        paras[anchor_idx].add_tracked_insertion(
                            f"\n{extra_text}", author=author
                        )

            # Extra old paragraphs → tracked deletions
            if len(old_chunk) > len(new_chunk):
                for k in range(pairs, len(old_chunk)):
                    del_idx = i1 + k
                    if del_idx < len(paras):
                        _delete_entire_paragraph(paras[del_idx], author)

        elif tag == "delete":
            # Old paragraphs removed
            for idx in range(i1, i2):
                if idx < len(paras):
                    _delete_entire_paragraph(paras[idx], author)

        elif tag == "insert":
            # New paragraphs added — insert as tracked insertion at anchor
            anchor_idx = max(0, i1 - 1)
            if anchor_idx < len(paras):
                for j in range(j1, j2):
                    paras[anchor_idx].add_tracked_insertion(
                        f"\n{new_texts[j]}", author=author
                    )

    _enable_markup_view(rdoc)
    return rdoc


def _apply_word_level_diff(
    para: RevisionParagraph,
    old_text: str,
    new_text: str,
    author: str,
) -> None:
    """Apply fine-grained word-level changes to a paragraph as tracked changes."""
    if old_text == new_text:
        return

    # Use the paragraph's replace_tracked for the simplest case
    # Find differing segments using SequenceMatcher on words
    old_words = old_text.split()
    new_words = new_text.split()

    sm = difflib.SequenceMatcher(None, old_words, new_words)

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue

        old_segment = " ".join(old_words[i1:i2])
        new_segment = " ".join(new_words[j1:j2])

        if tag == "replace" and old_segment and new_segment:
            para.replace_tracked(
                search_text=old_segment,
                replace_text=new_segment,
                author=author,
            )
        elif tag == "delete" and old_segment:
            # Find the position and mark as deleted
            text = old_text
            pos = text.find(old_segment)
            if pos >= 0:
                para.add_tracked_deletion(pos, pos + len(old_segment), author=author)
        elif tag == "insert" and new_segment:
            # Insert at the appropriate position
            # Find where to insert — after the previous equal block
            if i1 > 0:
                prev_word = old_words[i1 - 1]
                pos = old_text.find(prev_word)
                if pos >= 0:
                    insert_pos = pos + len(prev_word)
                    para.add_tracked_insertion(
                        f" {new_segment}", author=author
                    )
            else:
                para.add_tracked_insertion(
                    f"{new_segment} ", author=author
                )


def _delete_entire_paragraph(para: RevisionParagraph, author: str) -> None:
    """Mark all text in a paragraph as a tracked deletion."""
    text = ""
    if hasattr(para, "accepted_text"):
        text = para.accepted_text
    if not text and hasattr(para, "_paragraph"):
        text = para._paragraph.text
    if text:
        para.add_tracked_deletion(0, len(text), author=author)
