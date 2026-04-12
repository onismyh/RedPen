"""Wrapper around docx-revisions for applying tracked changes to Word documents."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
import hashlib

from lxml import etree
from docx.oxml.ns import qn
from docx_revisions import RevisionDocument, RevisionParagraph

from .format_preserver import copy_format_to_runs

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WNS = f"{{{WORD_NS}}}"


def _enable_markup_view(rdoc: RevisionDocument) -> None:
    """Ensure Word opens the document with tracked changes visible.

    Sets <w:revisionView> with markup=true in settings.xml so that
    Word defaults to 'All Markup' view instead of hiding revisions.
    """
    doc = rdoc.document
    settings_el = doc.settings.element

    # Remove existing revisionView if any
    for rv in settings_el.findall(f"{WNS}revisionView"):
        settings_el.remove(rv)

    # Add <w:revisionView w:markup="1" w:comments="1"
    #       w:insDel="1" w:formatting="1" />
    rv = etree.SubElement(settings_el, f"{WNS}revisionView")
    rv.set(f"{WNS}markup", "1")
    rv.set(f"{WNS}comments", "1")
    rv.set(f"{WNS}insDel", "1")
    rv.set(f"{WNS}formatting", "1")


@dataclass
class TextChange:
    """A single text change within a paragraph."""

    original: str
    revised: str
    reason: str = ""


@dataclass
class ParagraphEdit:
    """All changes for one paragraph, identified by index."""

    paragraph_index: int
    changes: list[TextChange]
    paragraph_anchor: str | None = None
    paragraph_text_hash: str | None = None


@dataclass
class ApplyStats:
    """Execution stats for applying edits."""

    proposed_changes: int = 0
    applied_changes: int = 0
    skipped_changes: int = 0
    not_found_changes: int = 0
    drifted_changes: int = 0
    applied_paragraph_indexes: set[int] | None = None


def compute_paragraph_text_hash(text: str) -> str:
    """Stable hash for paragraph text anchoring."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _paragraph_drift_warning(edit: ParagraphEdit, para_text: str) -> dict | None:
    """Return a warning if optional anchor/hash metadata disagrees with live text."""
    if edit.paragraph_anchor and edit.paragraph_anchor not in para_text:
        return {
            "kind": "paragraph_anchor_mismatch",
            "paragraph_index": edit.paragraph_index,
            "paragraph_anchor": edit.paragraph_anchor[:80],
        }

    if edit.paragraph_text_hash and edit.paragraph_text_hash != compute_paragraph_text_hash(para_text):
        return {
            "kind": "paragraph_text_hash_mismatch",
            "paragraph_index": edit.paragraph_index,
            "expected_hash": edit.paragraph_text_hash,
            "actual_hash": compute_paragraph_text_hash(para_text),
        }

    return None


def _replace_tracked_with_format(para, search_text: str, replace_text: str, author: str) -> int:
    """Replace text with tracked changes, preserving the original run formatting.

    After replace_tracked creates del/ins runs, copies rPr from the
    surrounding (before/after) runs so the replacement inherits the
    original font, bold, italic, color, etc.
    """
    count = para.replace_tracked(
        search_text=search_text,
        replace_text=replace_text,
        author=author,
    )

    if count > 0:
        # Copy formatting from surrounding runs to new del/ins runs
        copy_format_to_runs(para._p, para)

    return count


def apply_tracked_changes(
    doc_path: str,
    edits: list[ParagraphEdit],
    author: str = "AI Reviewer",
) -> RevisionDocument:
    """Open a document and apply a list of edits as tracked changes.

    Each edit locates text in a specific paragraph and replaces it
    using the revision tracking API (produces <w:del> + <w:ins>).
    Original formatting (bold, italic, font, color) is preserved.
    """
    rdoc, _, _, _ = apply_tracked_changes_checked(doc_path, edits, author=author)
    return rdoc



def apply_tracked_changes_checked(
    doc_path: str,
    edits: list[ParagraphEdit],
    author: str = "AI Reviewer",
) -> tuple[RevisionDocument, list[dict], ApplyStats, list[ParagraphEdit]]:
    """Apply tracked changes with paragraph drift detection but no protected-span filtering."""
    rdoc = RevisionDocument(doc_path)
    warnings: list[dict] = []
    applied_changes_by_paragraph: dict[int, list[TextChange]] = {}
    stats = ApplyStats(
        proposed_changes=sum(len(edit.changes) for edit in edits),
        applied_paragraph_indexes=set(),
    )

    for edit in edits:
        idx = edit.paragraph_index
        if idx < 0 or idx >= len(rdoc.paragraphs):
            stats.not_found_changes += len(edit.changes)
            continue

        para = rdoc.paragraphs[idx]
        para_text = para.accepted_text
        drift_warning = _paragraph_drift_warning(edit, para_text)
        if drift_warning is not None:
            warnings.append(drift_warning)
            stats.drifted_changes += len(edit.changes)
            continue

        for change in edit.changes:
            if change.original == change.revised:
                continue
            applied = _replace_tracked_with_format(
                para,
                search_text=change.original,
                replace_text=change.revised,
                author=author,
            )
            if applied > 0:
                stats.applied_changes += applied
                assert stats.applied_paragraph_indexes is not None
                stats.applied_paragraph_indexes.add(idx)
                applied_changes_by_paragraph.setdefault(idx, []).append(change)
            else:
                stats.not_found_changes += 1

    _enable_markup_view(rdoc)
    applied_edits = [
        ParagraphEdit(
            paragraph_index=idx,
            changes=changes,
            paragraph_anchor=next((e.paragraph_anchor for e in edits if e.paragraph_index == idx), None),
            paragraph_text_hash=next((e.paragraph_text_hash for e in edits if e.paragraph_index == idx), None),
        )
        for idx, changes in sorted(applied_changes_by_paragraph.items())
        if changes
    ]
    return rdoc, warnings, stats, applied_edits



def find_and_replace_tracked(
    doc_path: str,
    search_text: str,
    replace_text: str,
    author: str = "AI Reviewer",
) -> tuple[RevisionDocument, int]:
    """Document-wide find-and-replace with revision tracking.

    Preserves original formatting (bold, italic, font, color) on replacements.
    Returns (document, count_of_replacements).
    """
    rdoc = RevisionDocument(doc_path)
    count = 0

    # Body paragraphs
    for para in rdoc.paragraphs:
        count += _replace_tracked_with_format(para, search_text, replace_text, author)

    # Tables
    for table in rdoc.document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    rp = RevisionParagraph.from_paragraph(p)
                    count += _replace_tracked_with_format(rp, search_text, replace_text, author)

    _enable_markup_view(rdoc)
    return rdoc, count


def apply_tracked_changes_protected(
    doc_path: str,
    edits: list[ParagraphEdit],
    author: str = "AI Reviewer",
) -> tuple[RevisionDocument, list[dict], ApplyStats, list[ParagraphEdit]]:
    """Apply tracked changes with protection-aware filtering.

    Before applying each change, checks whether the original or revised text
    would modify a protected span (citation, formula, figure ref, etc.).
    Changes that alter protected spans are skipped and logged.

    Returns (document, list_of_skipped_warnings, execution_stats, applied_edits).
    """
    from .academic import find_protected_spans

    rdoc = RevisionDocument(doc_path)
    warnings: list[dict] = []
    applied_changes_by_paragraph: dict[int, list[TextChange]] = {}
    stats = ApplyStats(
        proposed_changes=sum(len(edit.changes) for edit in edits),
        applied_paragraph_indexes=set(),
    )

    for edit in edits:
        idx = edit.paragraph_index
        if idx < 0 or idx >= len(rdoc.paragraphs):
            stats.not_found_changes += len(edit.changes)
            continue

        para = rdoc.paragraphs[idx]
        para_text = para.accepted_text
        drift_warning = _paragraph_drift_warning(edit, para_text)
        if drift_warning is not None:
            warnings.append(drift_warning)
            stats.drifted_changes += len(edit.changes)
            continue

        protected = find_protected_spans(para_text)

        for change in edit.changes:
            if change.original == change.revised:
                continue

            # Check if this change would modify a protected span
            skip = False
            if protected:
                orig_pos = para_text.find(change.original)
                if orig_pos >= 0:
                    orig_end = orig_pos + len(change.original)
                    for span in protected:
                        # If the change region overlaps a protected span
                        if orig_pos < span.end and orig_end > span.start:
                            # Verify the protected text is preserved in revised
                            if span.text not in change.revised:
                                skip = True
                                warnings.append({
                                    "paragraph_index": idx,
                                    "kind": "protected_span_modified",
                                    "protected_text": span.text,
                                    "protected_kind": span.kind,
                                    "original": change.original[:80],
                                    "revised": change.revised[:80],
                                })
                                break

            if not skip:
                applied = _replace_tracked_with_format(
                    para,
                    search_text=change.original,
                    replace_text=change.revised,
                    author=author,
                )
                if applied > 0:
                    stats.applied_changes += applied
                    assert stats.applied_paragraph_indexes is not None
                    stats.applied_paragraph_indexes.add(idx)
                    applied_changes_by_paragraph.setdefault(idx, []).append(change)
                else:
                    stats.not_found_changes += 1
            else:
                stats.skipped_changes += 1

    _enable_markup_view(rdoc)
    applied_edits = [
        ParagraphEdit(paragraph_index=idx, changes=changes)
        for idx, changes in sorted(applied_changes_by_paragraph.items())
        if changes
    ]
    return rdoc, warnings, stats, applied_edits


def accept_all(doc_path: str) -> RevisionDocument:
    """Accept all tracked changes in a document."""
    rdoc = RevisionDocument(doc_path)
    rdoc.accept_all()
    return rdoc


def reject_all(doc_path: str) -> RevisionDocument:
    """Reject all tracked changes in a document."""
    rdoc = RevisionDocument(doc_path)
    rdoc.reject_all()
    return rdoc


def replace_cross_paragraph(
    doc_path: str,
    search_text: str,
    replace_text: str,
    author: str = "AI Reviewer",
) -> tuple[RevisionDocument, int]:
    """Find and replace across the entire document, including matches spanning paragraphs.

    Strategy:
    1. Join all paragraph texts with a unique separator.
    2. Find all match positions in the joined text.
    3. Map each match back to paragraph indices and within-paragraph offsets.
    4. Apply tracked replacements, handling cross-paragraph spans by
       deleting the spanned text in each paragraph and inserting the
       replacement at the first paragraph's match position.

    Returns (document, count_of_replacements).
    """
    SEP = "\x00"  # null byte as paragraph separator (won't appear in real text)

    rdoc = RevisionDocument(doc_path)
    paras = rdoc.paragraphs

    # Build joined text and track paragraph boundaries
    joined_parts: list[str] = []
    para_offsets: list[tuple[int, int]] = []  # (para_index, start_offset_in_joined)
    offset = 0
    for i, para in enumerate(paras):
        para_offsets.append((i, offset))
        text = para.accepted_text
        joined_parts.append(text)
        offset += len(text) + len(SEP)  # +1 for separator

    joined_text = SEP.join(joined_parts)

    # Find all matches
    count = 0
    search_len = len(search_text)
    start = 0
    while True:
        idx = joined_text.find(search_text, start)
        if idx == -1:
            break

        # Map match position to paragraph range
        match_end = idx + search_len

        # Find which paragraphs are involved
        start_para_idx = 0
        start_offset_in_para = 0
        end_para_idx = 0
        end_offset_in_para = 0

        for pi, (para_i, para_start) in enumerate(para_offsets):
            para_len = len(paras[para_i].accepted_text)
            para_end = para_start + para_len
            # SEP is at para_end
            sep_pos = para_end

            if idx >= para_start and idx <= sep_pos:
                start_para_idx = para_i
                start_offset_in_para = idx - para_start
                # Clamp to paragraph length
                start_offset_in_para = min(start_offset_in_para, para_len)

            if match_end >= para_start and match_end <= sep_pos:
                end_para_idx = para_i
                end_offset_in_para = match_end - para_start
                end_offset_in_para = min(end_offset_in_para, para_len)
                break

        # Apply replacements
        if start_para_idx == end_para_idx:
            # Single paragraph match — straightforward
            para = paras[start_para_idx]
            _replace_tracked_with_format(para, search_text, replace_text, author)
        else:
            # Cross-paragraph match
            # Delete the matched text from each involved paragraph
            # Insert replacement at the start position

            # Part 1: first paragraph — delete from start_offset to end
            first_para = paras[start_para_idx]
            first_para_text = first_para.accepted_text
            if start_offset_in_para < len(first_para_text):
                deleted_first = first_para_text[start_offset_in_para:]
                _delete_tracked_at(first_para, start_offset_in_para, len(first_para_text), author)

            # Part 2: middle paragraphs — delete entire text
            for pi in range(start_para_idx + 1, end_para_idx):
                mid_para = paras[pi]
                mid_text = mid_para.accepted_text
                if mid_text:
                    _delete_tracked_at(mid_para, 0, len(mid_text), author)

            # Part 3: last paragraph — delete from start to end_offset
            last_para = paras[end_para_idx]
            last_para_text = last_para.accepted_text
            if end_offset_in_para > 0 and end_offset_in_para <= len(last_para_text):
                _delete_tracked_at(last_para, 0, end_offset_in_para, author)

            # Part 4: insert replacement text at the first paragraph's match position
            # The replacement goes as a tracked insertion after the deleted portion
            first_para = paras[start_para_idx]
            first_para_text = first_para.accepted_text
            if replace_text:
                # Insert at the position where the original match started
                if start_offset_in_para < len(first_para_text):
                    # Insert after the deletion point — add at the end of the remaining text
                    # Since replace_tracked handles insertion, we use add_tracked_insertion
                    first_para.add_tracked_insertion(replace_text, author=author)
                else:
                    # Match was at the end of the paragraph (or empty para)
                    first_para.add_tracked_insertion(replace_text, author=author)

        count += 1
        start = idx + search_len

    _enable_markup_view(rdoc)
    return rdoc, count


def _delete_tracked_at(para, start: int, end: int, author: str) -> None:
    """Add a tracked deletion for text at [start, end) in a paragraph."""
    try:
        para.add_tracked_deletion(start, end, author=author)
    except (ValueError, IndexError):
        pass  # Silently skip invalid ranges
