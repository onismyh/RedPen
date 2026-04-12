"""RedPen CLI — Word track changes toolkit for AI agents."""

from __future__ import annotations

import io
import json as json_mod
import logging
import sys
from pathlib import Path

import typer
from rich.console import Console
from rich.table import Table
from rich.text import Text

app = typer.Typer(
    name="redpen",
    help="Word track changes toolkit for AI agents (Claude Code / Codex / OpenCode).",
    no_args_is_help=True,
)
console = Console()
logger = logging.getLogger("redpen")

from .recipes import app as recipe_app

app.add_typer(recipe_app, name="recipe", help="Generate task-oriented revision scaffolds.")


def _resolve_output(input_path: str, output: str | None) -> str:
    """If -o is given use it, otherwise overwrite the input file (in-place)."""
    if output:
        from pathlib import Path
        out = Path(output).resolve()
        inp = Path(input_path).resolve()
        if out == inp:
            logger.warning("Output path is the same as input path. File will be overwritten.")
    return output if output else input_path


# ---------------------------------------------------------------------------
# redpen read
# ---------------------------------------------------------------------------
@app.command()
def read(
    input_file: str = typer.Argument(..., help=".docx file to read"),
    plain: bool = typer.Option(False, "--plain", help="Plain text instead of JSON"),
    meta: bool = typer.Option(False, "--meta", help="Include paragraph style and run formatting metadata"),
) -> None:
    """Extract document paragraphs as JSON (default) or plain text.

    If the document already has tracked changes (from a previous round),
    the text shown is the "accepted" version — i.e. what the document
    looks like after accepting all prior revisions.  This lets agents
    make further edits on the latest state.

    With --meta, each paragraph includes:
      - style: paragraph style name (e.g. "Heading 1", "Normal")
      - is_heading: true if style contains "heading" or "title"
      - is_list: true if style contains "list" or "bullet"
      - runs: list of run objects with bold, italic, font_name
    """
    from docx.oxml.ns import qn
    from docx_revisions import RevisionDocument

    rdoc = RevisionDocument(input_file)
    paragraphs = []
    for i, para in enumerate(rdoc.paragraphs):
        text = para.accepted_text
        if not text.strip() and not meta:
            continue

        entry: dict = {"index": i, "text": text}

        if meta:
            # Paragraph style name
            style_name = ""
            try:
                style_name = para.style.name if para.style else ""
            except Exception:
                pass
            style_lower = style_name.lower()

            entry["style"] = style_name
            entry["is_heading"] = any(kw in style_lower for kw in ("heading", "title", "subtitle"))
            entry["is_list"] = any(kw in style_lower for kw in ("list", "bullet", "number"))

            # Run-level formatting
            runs_info = []
            for run in para.runs:
                if not run.text:
                    continue
                rpr = run._r.find(qn("w:rPr"))
                run_info: dict = {"text": run.text[:60]}  # truncate for readability

                if rpr is not None:
                    bold_el = rpr.find(qn("w:b"))
                    run_info["bold"] = bold_el is not None and bold_el.get(qn("w:val")) != "0"
                    italic_el = rpr.find(qn("w:i"))
                    run_info["italic"] = italic_el is not None and italic_el.get(qn("w:val")) != "0"
                    font_el = rpr.find(qn("w:rFonts"))
                    if font_el is not None:
                        run_info["font_name"] = font_el.get(qn("w:ascii")) or font_el.get(qn("w:hAnsi")) or ""
                    color_el = rpr.find(qn("w:color"))
                    if color_el is not None:
                        run_info["color"] = color_el.get(qn("w:val")) or ""
                else:
                    run_info["bold"] = False
                    run_info["italic"] = False
                    run_info["font_name"] = ""
                    run_info["color"] = ""

                runs_info.append(run_info)
            entry["runs"] = runs_info

        paragraphs.append(entry)

    has_changes = len(rdoc.track_changes) > 0

    if plain:
        for p in paragraphs:
            console.print(f"[{p['index']}] {p['text']}")
    else:
        print(json_mod.dumps(paragraphs, ensure_ascii=False, indent=2))

    if has_changes:
        Console(stderr=True).print(
            "[dim](document has existing tracked changes — "
            "text shown is the accepted/latest version)[/dim]"
        )


# ---------------------------------------------------------------------------
# redpen apply
# ---------------------------------------------------------------------------
@app.command()
def apply(
    input_file: str = typer.Argument(..., help="Input .docx file"),
    edits_json: str = typer.Argument(
        None, help="JSON string, @file path, or omit to read from stdin"
    ),
    output: str = typer.Option(None, "-o", "--output", help="Output file path"),
    author: str = typer.Option("Claude", "--author", help="Revision author name"),
    comment: bool = typer.Option(True, "--comment/--no-comment", help="Add comments from reason fields"),
) -> None:
    """Apply edits as Word tracked changes from JSON.

    Accepts JSON as:
      - Inline argument:  redpen apply doc.docx '[...]'
      - File reference:   redpen apply doc.docx @edits.json
      - Stdin pipe:       echo '[...]' | redpen apply doc.docx

    JSON format:

        [{"paragraph_index": 0, "changes": [
            {"original": "old", "revised": "new", "reason": "why"}
        ]}]
    """
    from .revision_writer import apply_tracked_changes, ParagraphEdit, TextChange
    from .comment_writer import add_comments_to_edits

    # Load JSON from argument, @file, or stdin
    if edits_json is None:
        if sys.stdin.isatty():
            console.print("[red]Error:[/red] No edits provided. Pass JSON, @file, or pipe via stdin.")
            raise typer.Exit(1)
        # Explicit UTF-8 to avoid locale-dependent encoding errors
        raw = io.TextIOWrapper(sys.stdin.buffer, encoding="utf-8").read()
    elif edits_json.startswith("@"):
        with open(edits_json[1:], "r", encoding="utf-8") as f:
            raw = f.read()
    else:
        raw = edits_json

    data = json_mod.loads(raw)
    edits = []
    for item in data:
        changes = [
            TextChange(
                original=c["original"],
                revised=c["revised"],
                reason=c.get("reason", ""),
            )
            for c in item.get("changes", [])
            if c.get("original") != c.get("revised")
        ]
        if changes:
            edits.append(ParagraphEdit(
                paragraph_index=item["paragraph_index"],
                changes=changes,
            ))

    if not edits:
        console.print("[yellow]No effective changes in the provided JSON.[/yellow]")
        raise typer.Exit(0)

    out_path = _resolve_output(input_file, output)
    rdoc = apply_tracked_changes(input_file, edits, author=author)

    if comment:
        doc = rdoc.document
        paragraphs = list(doc.paragraphs)
        add_comments_to_edits(doc, paragraphs, edits, author=author)

    rdoc.save(out_path)

    total = sum(len(e.changes) for e in edits)
    console.print(f"Applied {total} tracked changes across {len(edits)} paragraphs -> {out_path}")


# ---------------------------------------------------------------------------
# redpen replace
# ---------------------------------------------------------------------------
@app.command()
def replace(
    input_file: str = typer.Argument(..., help="Input .docx file"),
    search: str = typer.Argument(..., help="Text to find"),
    replacement: str = typer.Argument(..., help="Replacement text"),
    output: str = typer.Option(None, "-o", "--output", help="Output file path"),
    author: str = typer.Option("Claude", "--author", help="Revision author name"),
    cross_para: bool = typer.Option(False, "--cross-para", "-x", help="Allow matches spanning paragraph boundaries"),
) -> None:
    """Find and replace with revision tracking.

    Without --cross-para, only matches within a single paragraph are found.
    With --cross-para, matches that span across paragraphs are also replaced.
    """
    from .revision_writer import find_and_replace_tracked

    if cross_para:
        from .revision_writer import replace_cross_paragraph
        out_path = _resolve_output(input_file, output)
        rdoc, count = replace_cross_paragraph(input_file, search, replacement, author=author)
    else:
        out_path = _resolve_output(input_file, output)
        rdoc, count = find_and_replace_tracked(input_file, search, replacement, author=author)

    if count == 0:
        console.print(f"[yellow]Text not found:[/yellow] \"{search}\"")
        raise typer.Exit(1)

    rdoc.save(out_path)
    console.print(f"Replaced {count} occurrence(s) -> {out_path}")


# ---------------------------------------------------------------------------
# redpen diff
# ---------------------------------------------------------------------------
@app.command()
def diff(
    old_file: str = typer.Argument(..., help="Original .docx file"),
    new_file: str = typer.Argument(..., help="Modified .docx file"),
    output: str = typer.Option(None, "-o", "--output", help="Output file path"),
    author: str = typer.Option("Claude", "--author", help="Revision author name"),
) -> None:
    """Compare two documents, output a revision-tracked docx."""
    from .differ import diff_documents

    out_path = _resolve_output(old_file, output)
    rdoc = diff_documents(old_file, new_file, author=author)
    rdoc.save(out_path)
    console.print(f"Diff complete -> {out_path}")


# ---------------------------------------------------------------------------
# redpen show
# ---------------------------------------------------------------------------
@app.command()
def show(
    input_file: str = typer.Argument(..., help=".docx file with tracked changes"),
    json_output: bool = typer.Option(False, "--json", "-j", help="Output as JSON"),
) -> None:
    """Display tracked changes (table or JSON)."""
    from docx_revisions import RevisionDocument, TrackedInsertion, TrackedDeletion

    rdoc = RevisionDocument(input_file)
    changes = rdoc.track_changes

    if not changes:
        if json_output:
            print("[]")
        else:
            console.print("[yellow]No tracked changes found.[/yellow]")
        return

    if json_output:
        items = []
        for change in changes:
            items.append({
                "type": "insert" if isinstance(change, TrackedInsertion) else "delete",
                "author": change.author or "",
                "text": change.text,
            })
        print(json_mod.dumps(items, ensure_ascii=False, indent=2))
    else:
        table = Table(title=f"Tracked Changes - {Path(input_file).name}", show_lines=True)
        table.add_column("#", style="dim", width=4)
        table.add_column("Type", width=8)
        table.add_column("Author", width=14)
        table.add_column("Text", min_width=40)

        for i, change in enumerate(changes, 1):
            if isinstance(change, TrackedInsertion):
                change_type = Text("INSERT", style="green bold")
            elif isinstance(change, TrackedDeletion):
                change_type = Text("DELETE", style="red bold")
            else:
                change_type = Text("OTHER", style="yellow")

            text = change.text if change.text else ""
            display = text[:120] + ("..." if len(text) > 120 else "")
            table.add_row(str(i), change_type, change.author or "-", display)

        console.print(table)
        console.print(f"\nTotal: {len(changes)} tracked changes")


# ---------------------------------------------------------------------------
# redpen accept
# ---------------------------------------------------------------------------
@app.command()
def accept(
    input_file: str = typer.Argument(..., help=".docx file with tracked changes"),
    output: str = typer.Option(None, "-o", "--output", help="Output file path"),
) -> None:
    """Accept all tracked changes, produce a clean document."""
    from .revision_writer import accept_all

    out_path = _resolve_output(input_file, output)
    rdoc = accept_all(input_file)
    rdoc.save(out_path)
    console.print(f"All changes accepted -> {out_path}")


# ---------------------------------------------------------------------------
# redpen reject
# ---------------------------------------------------------------------------
@app.command()
def reject(
    input_file: str = typer.Argument(..., help=".docx file with tracked changes"),
    output: str = typer.Option(None, "-o", "--output", help="Output file path"),
) -> None:
    """Reject all tracked changes, restore original text."""
    from .revision_writer import reject_all

    out_path = _resolve_output(input_file, output)
    rdoc = reject_all(input_file)
    rdoc.save(out_path)
    console.print(f"All changes rejected -> {out_path}")


# ---------------------------------------------------------------------------
# redpen review  (academic review scaffold)
# ---------------------------------------------------------------------------
@app.command()
def review(
    input_file: str = typer.Argument(..., help=".docx file to review"),
    mode: str = typer.Option("proofread", "--mode", "-m", help="Review mode: proofread | academic-polish | reviewer"),
    lang: str = typer.Option(None, "--lang", "-l", help="Comment language: zh | en (default from config)"),
    run: bool = typer.Option(False, "--run", help="Actually call Claude to generate edits"),
    output: str = typer.Option(None, "-o", "--output", help="Output file path (used with --run, no --json)"),
    model: str = typer.Option(None, "--model", help="Claude model override"),
    json_output: bool = typer.Option(False, "--json", "-j", help="Output review plan as JSON"),
) -> None:
    """Generate a structured review plan for an academic paper.

    Without --run: produces a deterministic scaffold (no AI calls).
    With --run --json: calls Claude and prints generated edits as JSON.
    With --run: calls Claude, applies edits as tracked changes, saves to -o or <stem>.reviewed.docx.
    """
    from .config import load_config
    from .review import generate_review_plan

    if lang is None:
        lang = load_config().comment_language

    paragraphs = _read_paragraphs_with_meta(input_file)

    if run:
        import subprocess as _sp
        from .review import generate_edits_with_claude

        try:
            edits = generate_edits_with_claude(paragraphs, mode=mode, lang=lang, model=model)
        except _sp.TimeoutExpired:
            console.print("[red]Error:[/red] Claude timed out while generating edits.")
            raise typer.Exit(1)
        except _sp.CalledProcessError as exc:
            stderr = exc.stderr or ""
            console.print(f"[red]Error:[/red] Claude process failed: {stderr}")
            raise typer.Exit(1)

        if json_output:
            print(json_mod.dumps(edits, ensure_ascii=False, indent=2))
            return

        # Apply edits as tracked changes
        from .revision_writer import apply_tracked_changes, ParagraphEdit, TextChange
        from .comment_writer import add_comments_to_edits

        parsed_edits = []
        for item in edits:
            changes = [
                TextChange(
                    original=c["original"],
                    revised=c["revised"],
                    reason=c.get("reason", ""),
                )
                for c in item.get("changes", [])
                if c.get("original") != c.get("revised")
            ]
            if changes:
                parsed_edits.append(ParagraphEdit(
                    paragraph_index=item["paragraph_index"],
                    changes=changes,
                ))

        if not parsed_edits:
            console.print("[yellow]No effective changes generated.[/yellow]")
            raise typer.Exit(0)

        out_path = output or str(Path(input_file).with_suffix("").with_name(
            Path(input_file).stem + ".reviewed.docx"
        ))
        rdoc = apply_tracked_changes(input_file, parsed_edits, author="Claude")
        doc = rdoc.document
        add_comments_to_edits(doc, list(doc.paragraphs), parsed_edits, author="Claude")
        rdoc.save(out_path)

        total = sum(len(e.changes) for e in parsed_edits)

        # --- Generate clean docx (accept all tracked changes) ---
        out_p = Path(out_path)
        base = out_p.with_suffix("")  # strip .docx
        clean_path = base.parent / (base.name + ".clean.docx")
        report_path = base.parent / (base.name + ".report.json")

        from .revision_writer import accept_all
        clean_rdoc = accept_all(str(out_p))
        clean_rdoc.save(str(clean_path))

        # --- Generate report JSON ---
        report = {
            "mode": mode,
            "lang": lang,
            "source_file": str(Path(input_file).resolve()),
            "reviewed_file": str(out_p.resolve()),
            "clean_file": str(clean_path.resolve()),
            "report_file": str(report_path.resolve()),
            "change_count": total,
            "paragraph_count": len(paragraphs),
            "paragraphs_changed": len(parsed_edits),
            "model": model or "",
        }
        report_path.write_text(
            json_mod.dumps(report, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        console.print(f"Applied {total} tracked changes across {len(parsed_edits)} paragraphs -> {out_path}")
        console.print("Review complete.")
        return

    try:
        plan = generate_review_plan(paragraphs, mode=mode, lang=lang)
    except ValueError:
        console.print(f"[red]Error:[/red] Unknown review mode '{mode}'. Use: proofread, academic-polish, reviewer")
        raise typer.Exit(1)

    if json_output:
        print(plan.to_json())
    else:
        console.print(f"[bold]Review plan[/bold] — mode={plan.mode}, lang={plan.lang}")
        console.print(f"Paragraphs: {plan.total_paragraphs} total, {plan.reviewable_paragraphs} reviewable, {plan.skipped_paragraphs} skipped")
        console.print(f"Sections: {', '.join(plan.sections_found)}")
        console.print(f"Review items: {len(plan.items)}")
        if plan.items:
            table = Table(show_lines=False)
            table.add_column("Para#", style="dim", width=6)
            table.add_column("Section", width=16)
            table.add_column("Category", width=20)
            table.add_column("Severity", width=8)
            for item in plan.items[:30]:  # cap display
                table.add_row(str(item.paragraph_index), item.section, item.category, item.severity)
            console.print(table)
            if len(plan.items) > 30:
                console.print(f"[dim]... and {len(plan.items) - 30} more items[/dim]")


# ---------------------------------------------------------------------------
# redpen inspect  (academic structure inspection)
# ---------------------------------------------------------------------------
@app.command()
def inspect(
    input_file: str = typer.Argument(..., help=".docx file to inspect"),
    json_output: bool = typer.Option(False, "--json", "-j", help="Output as JSON"),
) -> None:
    """Inspect academic document structure: sections, headings, protected regions.

    Shows which sections were detected, which paragraphs are protected
    (citations, formulas, references), and a summary of the document layout.
    """
    from .academic import build_academic_structure

    paragraphs = _read_paragraphs_with_meta(input_file)
    structure = build_academic_structure(paragraphs)

    if json_output:
        items = []
        for ap in structure:
            entry: dict = {
                "index": ap.index,
                "section": ap.section.value,
                "is_heading": ap.is_heading,
                "is_protected": ap.is_protected,
                "text_preview": ap.text[:80] + ("..." if len(ap.text) > 80 else ""),
            }
            if ap.protected_spans:
                entry["protected_spans"] = [
                    {"kind": s.kind, "text": s.text} for s in ap.protected_spans
                ]
            items.append(entry)
        print(json_mod.dumps(items, ensure_ascii=False, indent=2))
    else:
        table = Table(title=f"Document Structure — {Path(input_file).name}", show_lines=False)
        table.add_column("#", style="dim", width=5)
        table.add_column("Section", width=16)
        table.add_column("H?", width=3)
        table.add_column("P?", width=3)
        table.add_column("Protected", width=12)
        table.add_column("Text", min_width=40)

        for ap in structure:
            if not ap.text.strip():
                continue
            h_mark = "H" if ap.is_heading else ""
            p_mark = "P" if ap.is_protected else ""
            pspan_summary = ", ".join(sorted({s.kind for s in ap.protected_spans})) if ap.protected_spans else ""
            preview = ap.text[:60] + ("..." if len(ap.text) > 60 else "")
            table.add_row(str(ap.index), ap.section.value, h_mark, p_mark, pspan_summary, preview)

        console.print(table)
        console.print(f"\nTotal: {len(structure)} paragraphs")

        # Section summary
        from collections import Counter
        sec_counts = Counter(ap.section.value for ap in structure if not ap.is_heading)
        console.print("\n[bold]Sections:[/bold]")
        for sec, cnt in sec_counts.most_common():
            console.print(f"  {sec}: {cnt} paragraphs")


# ---------------------------------------------------------------------------
# redpen check  (lightweight rule-based scan)
# ---------------------------------------------------------------------------
@app.command()
def check(
    input_file: str = typer.Argument(..., help=".docx file to check"),
    lang: str = typer.Option(None, "--lang", "-l", help="Comment language: zh | en (default from config)"),
    json_output: bool = typer.Option(False, "--json", "-j", help="Output as JSON"),
) -> None:
    """Run lightweight, rule-based checks on a document.

    Finds protected regions (citations, formulas, refs), overly long
    sentences, and other quick heuristics. No AI/LLM needed.
    """
    from .config import load_config
    from .review import quick_check

    if lang is None:
        lang = load_config().comment_language

    paragraphs = _read_paragraphs_for_check(input_file)
    results = quick_check(paragraphs, lang=lang)

    if json_output:
        items = [{"paragraph_index": r.paragraph_index, "kind": r.kind, "message": r.message} for r in results]
        print(json_mod.dumps(items, ensure_ascii=False, indent=2))
    else:
        if not results:
            console.print("[green]No issues found.[/green]")
            return
        table = Table(show_lines=False)
        table.add_column("Para#", style="dim", width=6)
        table.add_column("Kind", width=18)
        table.add_column("Message", min_width=40)
        for r in results:
            table.add_row(str(r.paragraph_index), r.kind, r.message[:100])
        console.print(table)
        console.print(f"\nTotal: {len(results)} findings")


# ---------------------------------------------------------------------------
# Shared helpers for review/inspect/check
# ---------------------------------------------------------------------------

def _read_paragraphs_with_meta(doc_path: str) -> list[dict]:
    """Read paragraphs with metadata (style, heading/list flags) for academic analysis."""
    from docx_revisions import RevisionDocument

    rdoc = RevisionDocument(doc_path)
    paragraphs = []
    for i, para in enumerate(rdoc.paragraphs):
        text = para.accepted_text
        style_name = ""
        try:
            style_name = para.style.name if para.style else ""
        except Exception:
            pass
        style_lower = style_name.lower()
        paragraphs.append({
            "index": i,
            "text": text,
            "style": style_name,
            "is_heading": any(kw in style_lower for kw in ("heading", "title", "subtitle")),
            "is_list": any(kw in style_lower for kw in ("list", "bullet", "number")),
        })
    return paragraphs


def _read_paragraphs_for_check(doc_path: str) -> list[dict]:
    """Read paragraphs (index + text only) for lightweight checks."""
    from docx_revisions import RevisionDocument

    rdoc = RevisionDocument(doc_path)
    return [{"index": i, "text": para.accepted_text} for i, para in enumerate(rdoc.paragraphs)]


if __name__ == "__main__":
    app()
