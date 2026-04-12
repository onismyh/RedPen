"""Task-oriented recipe scaffolds for common RedPen workflows."""

from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path

import typer
from docx import Document
from rich.console import Console

console = Console()
app = typer.Typer(help="Common revision workflows that generate starter scaffolds.")


@dataclass(frozen=True)
class RecipeSpec:
    name: str
    intent: str
    prompt: str
    change_reason: str
    description: str


RECIPES: dict[str, RecipeSpec] = {
    "proofread": RecipeSpec(
        name="proofread",
        intent="Fix grammar, punctuation, and wording while preserving meaning.",
        prompt=(
            "Proofread this Word document paragraph by paragraph. Keep the author's meaning, "
            "make minimal edits, and explain each change briefly in the reason field."
        ),
        change_reason="Fix grammar or improve clarity while preserving meaning.",
        description="Best for grammar cleanup, typo fixes, and light polishing.",
    ),
    "tighten": RecipeSpec(
        name="tighten",
        intent="Make writing more concise and executive without changing claims.",
        prompt=(
            "Tighten this document. Remove redundancy, shorten weak phrases, and keep the tone professional. "
            "Explain each reduction in the reason field."
        ),
        change_reason="Make the sentence more concise and direct.",
        description="Best for reducing fluff in reports, updates, and memos.",
    ),
    "reviewer": RecipeSpec(
        name="reviewer",
        intent="Act like a human reviewer using Track Changes and comments.",
        prompt=(
            "Review this document like a human reviewer. Use conservative revisions, preserve the author's voice, "
            "and write comments that explain concerns or recommendations."
        ),
        change_reason="Reviewer note explaining the suggested change.",
        description="Best for collaborative review where the author will Accept / Reject each edit.",
    ),
    "academic": RecipeSpec(
        name="academic",
        intent="Polish an English academic paper: fix grammar, improve clarity, and tighten prose while preserving technical meaning, citations, formulas, and references.",
        prompt=(
            "Review this academic paper paragraph by paragraph. "
            "Fix grammar errors, improve clarity and conciseness, and ensure consistent academic tone. "
            "Do NOT modify: citations like [1] or (Author, 2024), inline math ($...$), "
            "Figure/Table/Equation references (e.g. Fig. 1, Table 2, Eq. 3), URLs, or DOIs. "
            "Skip paragraphs in the References/Bibliography section entirely. "
            "Explain each change briefly in the reason field."
        ),
        change_reason="Academic polish: improve clarity or fix grammar while preserving technical meaning.",
        description="Best for polishing English academic papers before submission. Protects citations, formulas, and references.",
    ),
}


def _starter_edits(doc_path: Path, reason: str) -> list[dict]:
    doc = Document(str(doc_path))
    edits: list[dict] = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if len(text) < 12:
            continue
        edits.append(
            {
                "paragraph_index": idx,
                "changes": [
                    {
                        "original": text,
                        "revised": text,
                        "reason": reason,
                    }
                ],
            }
        )
        if len(edits) == 2:
            break
    return edits


def _build_payload(recipe: RecipeSpec, input_file: Path) -> dict:
    return {
        "recipe": recipe.name,
        "input_file": str(input_file),
        "description": recipe.description,
        "prompt": recipe.prompt,
        "suggested_command": f"redpen apply {input_file} @{recipe.name}.edits.json -o revised.docx",
        "starter_edits": _starter_edits(input_file, recipe.change_reason),
    }


@app.command()
def proofread(
    input_file: str = typer.Argument(..., help="Input .docx file"),
    json_output: bool = typer.Option(False, "--json", help="Output recipe scaffold as JSON"),
) -> None:
    """Generate a starter scaffold for proofreading."""
    payload = _build_payload(RECIPES["proofread"], Path(input_file).resolve())
    if json_output:
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return

    console.print("[bold]RedPen recipe: proofread[/bold]")
    console.print("Use Track Changes to fix grammar and clarity without rewriting the whole document.")
    console.print(f"Input: {payload['input_file']}")
    console.print(f"Prompt: {payload['prompt']}")
    console.print(f"Next: {payload['suggested_command']}")


@app.command()
def tighten(
    input_file: str = typer.Argument(..., help="Input .docx file"),
    json_output: bool = typer.Option(False, "--json", help="Output recipe scaffold as JSON"),
) -> None:
    """Generate a starter scaffold for concise revision."""
    payload = _build_payload(RECIPES["tighten"], Path(input_file).resolve())
    if json_output:
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return

    console.print("[bold]RedPen recipe: tighten[/bold]")
    console.print("Use Track Changes to compress wording without changing claims.")
    console.print(f"Input: {payload['input_file']}")
    console.print(f"Prompt: {payload['prompt']}")
    console.print(f"Next: {payload['suggested_command']}")


@app.command()
def reviewer(
    input_file: str = typer.Argument(..., help="Input .docx file"),
    json_output: bool = typer.Option(False, "--json", help="Output recipe scaffold as JSON"),
) -> None:
    """Generate a starter scaffold for reviewer-style feedback."""
    payload = _build_payload(RECIPES["reviewer"], Path(input_file).resolve())
    if json_output:
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return

    console.print("[bold]RedPen recipe: reviewer[/bold]")
    console.print("Review with native Word Track Changes so the author can Accept / Reject each edit.")
    console.print("This mode is best for editorial or collaborative reviewer workflows.")
    console.print(f"Input: {payload['input_file']}")
    console.print(f"Prompt: {payload['prompt']}")
    console.print(f"Next: {payload['suggested_command']}")


@app.command()
def academic(
    input_file: str = typer.Argument(..., help="Input .docx file"),
    comment_lang: str = typer.Option(None, "--comment-lang", help="Comment language: zh | en (default from config)"),
    json_output: bool = typer.Option(False, "--json", help="Output recipe scaffold as JSON"),
) -> None:
    """Generate a starter scaffold for academic paper polishing."""
    from .config import load_config

    if comment_lang is None:
        comment_lang = load_config().comment_language

    payload = _build_payload(RECIPES["academic"], Path(input_file).resolve())

    # Section-aware: detect and include academic sections in payload
    from .academic import build_academic_structure
    from docx_revisions import RevisionDocument

    rdoc = RevisionDocument(str(Path(input_file).resolve()))
    paragraphs = []
    for i, para in enumerate(rdoc.paragraphs):
        text = para.accepted_text
        style_name = ""
        try:
            style_name = para.style.name if para.style else ""
        except Exception:
            pass
        style_lower = style_name.lower()
        paragraphs.append({
            "index": i,
            "text": text,
            "style": style_name,
            "is_heading": any(kw in style_lower for kw in ("heading", "title", "subtitle")),
        })
    structure = build_academic_structure(paragraphs)
    payload["sections"] = sorted({ap.section.value for ap in structure})

    payload["comment_language"] = comment_lang
    if comment_lang == "zh":
        payload["prompt"] += " Write all change reasons and comments in Chinese (中文)."
    else:
        payload["prompt"] += " Write all change reasons and comments in English."

    if json_output:
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return

    console.print("[bold]RedPen recipe: academic[/bold]")
    console.print("Polish an English academic paper using Track Changes.")
    console.print("Citations, formulas, figure/table/equation references are protected.")
    console.print(f"Comment language: {comment_lang}")
    console.print(f"Input: {payload['input_file']}")
    console.print(f"Prompt: {payload['prompt']}")
    console.print(f"Next: {payload['suggested_command']}")
